"""Generate SPAR overhaul README as a Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SPAR_Overhaul_And_Updates.docx"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("SPAR & Quorum-CLI — Architecture Overhaul & Updates", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _para(doc, "Document version: July 2026")
    _para(doc, "Project: quorum-cli (SP Jain — AI in Finance, Group 3)")
    _para(doc, "Repository: https://github.com/GoldnDarkns/quorum-cli")
    doc.add_paragraph()

    _heading(doc, "1. Executive Summary")
    _para(
        doc,
        "This document describes all major updates made to the Quorum-CLI fork for the SPAR "
        "(Scenario Planning via Agentic Reasoning) research system. The central upgrade moves "
        "SPAR from event-analogue-heavy grounding to a transmission-channel-first evidence "
        "pipeline (Layer 0), followed by a live sequential multi-agent debate (Layer 1).",
    )
    doc.add_paragraph()

    _heading(doc, "2. What Changed — At a Glance")
    for item in [
        "New Layer 0 control pipeline (deterministic Python, no extra LLM calls)",
        "12 hardcoded financial transmission channels with explainable scoring",
        "Per-channel evidence retrieval replaces top-3 historical analogue stuffing",
        "Agent-specific evidence packets routed by domain",
        "SPAR expanded from 3 phases to 4 phases (Layer 0 added)",
        "Round 2 redesigned as live sequential debate — agents read and respond to each other",
        "Offline Ollama support tuned for RTX 3050 Ti (4 GB VRAM) laptops",
        "Model validation timeout fixes for slow local models",
        "JSON output schema updated: channel_assessment replaces analogue_assessment",
        "Terminal UI i18n updated for new SPAR phase names",
    ]:
        _bullet(doc, item)
    doc.add_paragraph()

    _heading(doc, "3. Architecture Overhaul — Transmission-Channel-First RAG")
    _heading(doc, "3.1 Problem with the Old Design", 2)
    _para(
        doc,
        "The original master_context.txt injected three fixed historical analogues "
        "(Kuwait 1990, Crimea 2014, Iraq War 2003) into every agent's prompt. All five "
        "specialists saw the same anchor events, which encouraged groupthink and shallow "
        "reasoning. The system asked 'which past wars look similar?' instead of 'which "
        "financial transmission mechanisms does this shock activate?'",
    )
    doc.add_paragraph()

    _heading(doc, "3.2 New Layer 0 Pipeline", 2)
    _para(doc, "File: src/quorum/methods/spar_layer0.py")
    _para(doc, "Flow:")
    _bullet(doc, "0.1 Regime Classifier — macro/market regime before shock evaluation")
    _bullet(doc, "0.2 Shock Parser — entities, event type, affected systems, time horizon")
    _bullet(doc, "0.3 Transmission Channel Library — 12 reusable financial mechanisms")
    _bullet(doc, "0.4 Channel Prioritizer — explainable scoring formula (0–100)")
    _bullet(doc, "0.5 RAG Query Builder — channel-specific retrieval queries")
    _bullet(doc, "0.6 Evidence Retriever — curated per-channel evidence corpus")
    _bullet(doc, "0.7 Evidence Sufficiency Scorer — flags thin evidence bases")
    _bullet(doc, "0.8 Evidence Packet Builder + Agent Router — domain-specific context")
    doc.add_paragraph()

    _heading(doc, "3.3 Channel Scoring Formula", 2)
    _para(
        doc,
        "Channel Activation Score = 0.30 × Event/entity match + 0.25 × Economic mechanism match "
        "+ 0.20 × Macro-regime relevance + 0.15 × Historical evidence availability "
        "+ 0.10 × Market/sector materiality",
    )
    _para(doc, "Priority tiers:")
    _bullet(doc, "75–100: Primary (5–7 evidence items)")
    _bullet(doc, "50–74: Secondary (2–3 evidence items)")
    _bullet(doc, "30–49: Watchlist (1 evidence item)")
    _bullet(doc, "<30: Inactive (ignored unless flagged later)")
    doc.add_paragraph()

    _heading(doc, "3.4 Twelve Transmission Channels", 2)
    channels = [
        "Geopolitical Risk Premium",
        "Energy / Commodity Price Shock",
        "Inflation Shock",
        "Monetary Policy Constraint",
        "Sanctions / Trade / Policy Shock",
        "Supply Chain Disruption",
        "Safe-Haven / FX Flow",
        "Credit / Financial Conditions",
        "Sector Earnings Exposure",
        "Consumer Sentiment / Behavioural Shock",
        "Cyber / Operational Disruption",
        "Defence Spending Repricing",
        "Relief Rally / Priced-In Shock Dampener",
    ]
    for ch in channels:
        _bullet(doc, ch)
    doc.add_paragraph()

    _heading(doc, "4. SPAR Debate Flow (Layer 1)")
    _heading(doc, "Phase 1 — Layer 0: Channel Prioritization", 2)
    _para(
        doc,
        "Runs before any agent speaks. Displays activated channels and scores in the terminal. "
        "No API keys required for this step — it is pure Python.",
    )
    _heading(doc, "Phase 2 — Round 1: Independent Domain Analysis", 2)
    _para(
        doc,
        "Each of five specialists (Political, Economic, Environmental, Social, Devil's Advocate) "
        "produces structured JSON with direction, magnitude estimates, confidence, transmission "
        "channels, and channel_assessment. Agents receive routed Layer 0 evidence packets.",
    )
    _heading(doc, "Phase 3 — Round 2: Live Sequential Debate", 2)
    _para(
        doc,
        "NEW: Agents no longer receive a silent JSON dump and reply in isolation. They speak "
        "in turn and see the full growing transcript — including other agents' live responses "
        "in the same round. Each agent must name another agent, cite their claim, agree or "
        "disagree with evidence, and speak in prose (not JSON). This mimics a real war-room panel.",
    )
    _heading(doc, "Phase 4 — Moderator Synthesis", 2)
    _para(
        doc,
        "Moderator receives Layer 0 state, Round 1 JSON, Round 2 live transcript, and produces "
        "consensus scenario + minority dissent with plausibility scoring based on channel consistency.",
    )
    doc.add_paragraph()

    _heading(doc, "5. Prompt & Schema Updates")
    _bullet(doc, "master_context.txt — removed fixed 3-analogue block; added channel-first instructions")
    _bullet(doc, "moderator.txt — plausibility scored on channel evidence consistency")
    _bullet(doc, "JSON schema — analogue_assessment replaced by channel_assessment")
    _bullet(doc, "Round 2 — prose live debate instead of structured JSON cross-examination")
    doc.add_paragraph()

    _heading(doc, "6. Offline / Local Model Setup")
    _para(doc, "Hardware profile used for tuning: Intel i7-11370H, 16 GB RAM, RTX 3050 Ti (4 GB VRAM).")
    _bullet(doc, "Ollama installed for local inference (no cloud API keys required)")
    _bullet(doc, "Recommended models: qwen2.5:7b + llama3.2:3b (or qwen2.5:3b for faster runs)")
    _bullet(doc, "QUORUM_EXECUTION_MODE=sequential — prevents VRAM contention on 4 GB GPU")
    _bullet(doc, "QUORUM_MODEL_TIMEOUT=300 — local models need longer cold-start time")
    _bullet(doc, "gpt-oss:20b available in Ollama but not recommended on this hardware (too large)")
    doc.add_paragraph()

    _heading(doc, "7. Technical Fixes Included in This Push")
    _bullet(doc, "models.py — Ollama validation uses QUORUM_MODEL_TIMEOUT from .env (was hardcoded 60s)")
    _bullet(doc, "ipc.py — Ollama model validation serialized to avoid parallel VRAM overload")
    _bullet(doc, "frontend — SPAR phase labels updated to 4 phases across all locales")
    _bullet(doc, "tests/test_spar_layer0.py — unit tests for channel prioritization pipeline")
    doc.add_paragraph()

    _heading(doc, "8. Files Added or Modified")
    _heading(doc, "New files", 2)
    _bullet(doc, "src/quorum/methods/spar_layer0.py — Layer 0 pipeline")
    _bullet(doc, "tests/test_spar_layer0.py — Layer 0 tests")
    _bullet(doc, "docs/SPAR_Overhaul_And_Updates.docx — this document")
    _heading(doc, "Modified files", 2)
    for f in [
        "src/quorum/methods/spar.py",
        "src/quorum/models.py",
        "src/quorum/ipc.py",
        "research/prompts/master_context.txt",
        "research/prompts/moderator.txt",
        "frontend/src/i18n/translations/*.ts",
        "frontend/src/utils/phases.ts",
    ]:
        _bullet(doc, f)
    doc.add_paragraph()

    _heading(doc, "9. How to Run")
    _para(doc, "1. Install Ollama and pull models: ollama pull qwen2.5:7b && ollama pull llama3.2:3b")
    _para(doc, "2. Configure .env (copy from .env.example): OLLAMA_BASE_URL, QUORUM_EXECUTION_MODE=sequential")
    _para(doc, '3. From project root in PowerShell: .\\quorum.bat')
    _para(doc, "4. Select both Ollama models via /models")
    _para(doc, "5. Run SPAR: /method spar  OR  /spar <your shock scenario question>")
    doc.add_paragraph()

    _heading(doc, "10. Future Roadmap (Not Yet Implemented)")
    _bullet(doc, "Full vector RAG with FRED, GPR index, EIA, sanctions databases (live APIs)")
    _bullet(doc, "Embeddings + vector store replacing curated evidence corpus")
    _bullet(doc, "Multi-round live debate (3+ speaking rounds with rebuttals)")
    _bullet(doc, "Cloud multi-model deployment for SPAR (5 different providers per agent)")
    _bullet(doc, "PDF export of full debate transcript with channel evidence appendix")
    _bullet(doc, "Deterministic shock parser for arbitrary novel scenarios beyond Ukraine pilot")
    doc.add_paragraph()

    _heading(doc, "11. Academic Rationale")
    _para(
        doc,
        "The overhaul aligns SPAR with the research thesis that financial shocks propagate "
        "through repeatable transmission channels rather than through surface-level event similarity. "
        "By making Layer 0 deterministic and auditable, reviewers can inspect why, for example, "
        "Energy Shock scored 92 without re-running an LLM. Layer 1 debate then tests whether "
        "domain-specialist models converge or dissent when given channel-routed evidence — a "
        "testable, scientifically meaningful experimental design.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
